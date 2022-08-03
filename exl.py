import string
from docx import Document
import pandas as pd
from docx2pdf import convert
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import time
data = pd.read_excel('Book1.xlsx')
# df = pd.DataFrame(data, columns= ['Email'])

user_name = ""
# getting the names and the emails
names = data['Name']
emails = data['Email']
salarys=data['Salry']
positions=data['Position']
# iterate through the records

# with open('sample.txt') as f:
#     lines = f.readlines()

# lines=str(lines)
# lines=lines.replace("['","")
# lines=lines.replace("\\n', '","\n")
# lines=lines.replace("']","")
# f.write("This line will be appended at the end")
# f.close()
 

# print(lines)





def sendMail(pdf_file_to_send,send_to,user_name):
    #start
    body = '''Hey %s,
    Please Check the message below
    '''%(user_name)
    # put your email here
    sender = 'learntoearn195@gmail.com'

    password = 'zdqvmgikgvxaxlsg'

    receiver = send_to
    
    #Setup the MIME
    message = MIMEMultipart()
    message['From'] = sender
    message['To'] = receiver
    message['Subject'] = 'Promotion Notification'
    
    message.attach(MIMEText(body, 'plain'))
    
    pdfname = pdf_file_to_send
    
    # open the file in bynary
    binary_pdf = open(pdfname, 'rb')
    
    payload = MIMEBase('application', 'octate-stream', Name=pdfname)
    # payload = MIMEBase('application', 'pdf', Name=pdfname)
    payload.set_payload((binary_pdf).read())
    
    # enconding the binary into base64
    encoders.encode_base64(payload)
    
    # add header with pdf name
    payload.add_header('Content-Decomposition', 'attachment', filename=pdfname)
    message.attach(payload)
    
    #use gmail with port
    session = smtplib.SMTP('smtp.gmail.com', 587)
    
    #enable security
    session.starttls()
    
    #login with mail_id and password
    session.login(sender, password)
    
    text = message.as_string()
    session.sendmail(sender, receiver, text)
    session.quit()
    print('Mail Sent')
    #end
def ch_var(name,salary,position):
        #start
        dic = {
            '[NAME]':name,
            '[POST]':position,
            '[SALARY]' : salary,
        }
        for p in sample_page.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                for key in dic.keys():
                    if key in text:
                        text=text.replace(key,dic[key])
                        inline[i].text = text
    #end    
for i in range(len(emails)):
    try:
        # for every record get the name and the email addresses
        sample_page = Document("tempalate.docx")
        n = names[i]
        email = emails[i]
        position= positions[i]
        sal= salarys[i]
        salary=str(sal)
        name=str(n)
    
        send_to=email
        # the message to be emailed
        ch_var(name,salary,position)
        

        # var1=lines.replace('NAME', name)
        # var2=var1.replace('POSITION', position)
        # var3=var2.replace('?salary', salary)

        # print(sample_page)

        # message = "Hey, "+name+ "  congratulation you are\npramoted to position\n" +position+ " and your salary is now "+str(salary)
        # create document object
        # doc = Document()


        # add text to the document
        # doc.add_paragraph(var3)

        # save document
        user_name = name
        # name = "output\\"+name
        sample_page.save("output\\"+name+'.docx')
        filename = "output\\"+name+".pdf"
        convert("output\\"+name+".docx", filename)

        print("pdf generation done")
        sendMail(filename,send_to,user_name)
        print("Mail sent: "+filename)
        print("waiting 30sec")
        time.sleep(30)
        print("cycle completed.************************")

    except Exception as e:
        print(e)

    # print(message)
print("process completed")

